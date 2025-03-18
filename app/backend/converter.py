import os
import logging
import customtkinter as ctk
from tkinter import filedialog, messagebox, Toplevel
from docx import Document
import subprocess
from googletrans import Translator

# Configure logging
logging.basicConfig(level=logging.DEBUG, filename='srt_to_word.log', filemode='w',
                    format='%(asctime)s - %(levelname)s - %(message)s')

# Translator instance
translator = Translator()

# Function to process SRT file
def srt_to_paragraphs(srt_file_path, char_limit=None):
    """Converts SRT to a single paragraph, respecting optional char limit."""
    paragraphs = []
    with open(srt_file_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()
    
    current_paragraph = []
    for line in lines:
        line = line.strip()
        if line.isdigit():  # Skip numbering
            continue
        if '-->' in line:  # Skip time codes
            continue
        if line == "":  # Empty line indicates end of a subtitle block
            if current_paragraph:
                paragraphs.append(" ".join(current_paragraph))
                current_paragraph = []
        else:
            current_paragraph.append(line)
    
    if current_paragraph:  # Add any remaining text
        paragraphs.append(" ".join(current_paragraph))
    
    # Combine all paragraphs into one big paragraph
    combined_paragraph = " ".join(paragraphs)
    
    # If a character limit is provided, break the text into paragraphs with the limit
    if char_limit:
        final_paragraphs = []
        current_text = ""
        for word in combined_paragraph.split():
            if len(current_text + " " + word) <= char_limit:
                current_text += " " + word
            else:
                final_paragraphs.append(current_text.strip())
                current_text = word
        if current_text:  # Add any remaining text
            final_paragraphs.append(current_text.strip())
        return "\n\n".join(final_paragraphs)
    
    return combined_paragraph

# Function to save to Word
def save_to_word(paragraphs, word_file_path, heading):
    try:
        doc = Document()
        if heading:
            doc.add_heading(heading, level=1)
        doc.add_paragraph(paragraphs)
        doc.save(word_file_path)
    except Exception as e:
        logging.error(f"Error saving Word document: {e}")
        raise e

# Function to translate SRT file
def translate_srt(srt_file_path, target_language="en"):
    """Translates the text content of an SRT file to the target language."""
    translated_lines = []
    with open(srt_file_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    for line in lines:
        line = line.strip()
        if line.isdigit() or '-->' in line or line == "":
            # Keep numbering, time codes, and empty lines unchanged
            translated_lines.append(line)
        else:
            # Translate subtitle text
            try:
                translated_text = translator.translate(line, dest=target_language).text
                translated_lines.append(translated_text)
            except Exception as e:
                logging.error(f"Error translating line: {line}. Error: {e}")
                translated_lines.append(line)  # Fallback to original line

    return "\n".join(translated_lines)

# Callback to handle file conversion
def convert_file():
    status_var.set("Converting...")
    root.update_idletasks()
    
    srt_files = input_file_paths.get()
    heading = heading_entry.get()
    char_limit = char_limit_entry.get()

    if not srt_files:
        messagebox.showerror("Error", "Please select at least one SRT file.")
        status_var.set("Ready")
        return

    srt_file_list = srt_files.split(';')
    selected_files_count.set(f"Selected Files: {len(srt_file_list)}")  # Show number of selected files
    
    # Check if all selected files exist
    for srt_file in srt_file_list:
        if not os.path.exists(srt_file):
            messagebox.showerror("Error", f"The selected SRT file does not exist: {srt_file}")
            status_var.set("Ready")
            return

    if char_limit:
        try:
            char_limit = int(char_limit)
        except ValueError:
            messagebox.showerror("Error", "Character limit must be an integer.")
            status_var.set("Ready")
            return

    output_dir = filedialog.askdirectory(title="Select Output Folder")
    if not output_dir:
        status_var.set("Ready")
        return

    processed_count = 0
    try:
        for srt_file in srt_file_list:
            paragraphs = srt_to_paragraphs(srt_file, char_limit)
            output_path = os.path.join(output_dir, os.path.basename(srt_file).replace(".srt", ".docx"))
            save_to_word(paragraphs, output_path, heading)
            processed_count += 1
            processed_files_count.set(f"Processed Files: {processed_count}")  # Update processed count
        messagebox.showinfo("Success", f"Files converted successfully!\nSaved in: {output_dir}")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")
    
    status_var.set("Ready")


# Callback to browse files
def browse_files():
    file_paths = filedialog.askopenfilename(
        filetypes=[("SRT Files", "*.srt")],
        title="Select SRT Files",
        multiple=True
    )
    if file_paths:
        input_file_paths.set(';'.join(file_paths))
        selected_files_count.set(f"Selected Files: {len(file_paths)}")  


# Modern GUI Design
ctk.set_appearance_mode("System")  
ctk.set_default_color_theme("blue")  

# Root window
root = ctk.CTk()
root.title("TextCraft Pro")
root.geometry("700x500")

# Status bar
status_var = ctk.StringVar(value="Ready")
status_label = ctk.CTkLabel(root, textvariable=status_var, fg_color="lightgrey", anchor="w", height=20)
status_label.pack(side="bottom", fill="x")

# Define the processed_files_count variable
processed_files_count = ctk.StringVar(value="Processed Files: 0")

# Input file section
frame = ctk.CTkFrame(master=root)
frame.pack(pady=20, padx=20, fill="both", expand=True)

title_label = ctk.CTkLabel(master=frame, text="SRT to Word Converter", font=("Arial", 24))
title_label.pack(pady=10)

file_label = ctk.CTkLabel(master=frame, text="Select SRT Files:")
file_label.pack(pady=5)
input_file_paths = ctk.StringVar()
file_entry = ctk.CTkEntry(master=frame, textvariable=input_file_paths, width=400)
file_entry.pack(pady=5)
file_button = ctk.CTkButton(master=frame, text="Browse", command=browse_files)
file_button.pack(pady=5)

# Label to show number of selected files
selected_files_count = ctk.StringVar(value="Selected Files: 0")
selected_files_label = ctk.CTkLabel(master=frame, textvariable=selected_files_count)
selected_files_label.pack(pady=5)

# Heading section
heading_label = ctk.CTkLabel(master=frame, text="Custom Heading (Optional):")
heading_label.pack(pady=5)
heading_entry = ctk.CTkEntry(master=frame, width=400)
heading_entry.pack(pady=5)

# Character limit section
char_limit_label = ctk.CTkLabel(master=frame, text="Max Characters Per Paragraph (Optional):")
char_limit_label.pack(pady=5)
char_limit_entry = ctk.CTkEntry(master=frame, width=400)
char_limit_entry.pack(pady=5)

# Convert button
convert_button = ctk.CTkButton(master=frame, text="Convert to Word", command=convert_file)
convert_button.pack(pady=20)

# Label to show number of processed files
processed_files_label = ctk.CTkLabel(master=frame, textvariable=processed_files_count)
processed_files_label.pack(pady=5)

# Run the app
root.mainloop()
